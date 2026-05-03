#!/usr/bin/env python3
"""Generate English résumé .docx from ``resume_data_en.RESUME`` (data-driven lists).

Uses python-docx only; narrative text lives in ``resume_data_en.py``.

Run with Lunaya Python image (``python-docx`` is listed in lunaya-api requirements)::

    docker compose -f docker-compose.dev.yml run --rm --no-deps \\
      -v /path/to/mohammadgholampour.github.io:/resume \\
      lunaya-api python /resume/scripts/build_resume_en_docx.py \\
        [--output /resume/Mohammad_Gholampour_Resume_EN_v2.docx]

Default ``--output`` is ``Mohammad_Gholampour_Resume_EN_v2.docx`` beside ``scripts/``
(i.e. the résumé repo root), so ``Mohammad_Gholampour_Resume_EN.docx`` is left untouched.
"""

from __future__ import annotations

import argparse
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

_SCRIPTS_DIR = Path(__file__).resolve().parent
_DEFAULT_RESUME_ROOT = _SCRIPTS_DIR.parent
_DEFAULT_OUTPUT = _DEFAULT_RESUME_ROOT / "Mohammad_Gholampour_Resume_EN_v2.docx"

# Word run shading fill (highlight) for KPI snippets — WCAG-friendly light amber
_KPI_SHADING_FILL = "FFF2CC"

_BODY_PT = Pt(10.5)
_SECTION_PT = Pt(11)
_NAME_PT = Pt(16)
_TITLE_PT = Pt(12)


def _load_resume(resume_root: Path) -> dict:
    root = resume_root.expanduser().resolve()
    if str(root) not in sys.path:
        sys.path.insert(0, str(root))
    from resume_data_en import RESUME  # noqa: PLC0415 — import after sys.path fix

    return RESUME


def _set_run_font_size(run, size_pt) -> None:
    run.font.size = size_pt


def _paragraph_spacing_after(paragraph, after_pt: float = 6) -> None:
    pf = paragraph.paragraph_format
    pf.space_after = Pt(after_pt)


def _paragraph_spacing_compact(paragraph) -> None:
    pf = paragraph.paragraph_format
    pf.space_after = Pt(2)
    pf.space_before = Pt(0)


def _run_shading_fill(run, fill_hex_rrggbb: str) -> None:
    r_pr = run._element.get_or_add_rPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill_hex_rrggbb)
    shd.set(qn("w:val"), "clear")
    r_pr.append(shd)


def _add_normal_run(paragraph, text: str, *, size=_BODY_PT) -> None:
    if not text:
        return
    r = paragraph.add_run(text)
    _set_run_font_size(r, size)
    if size == _BODY_PT:
        r.font.color.rgb = RGBColor(0x33, 0x33, 0x33)


def _add_bold_run(paragraph, text: str, *, size=_BODY_PT) -> None:
    if not text:
        return
    r = paragraph.add_run(text)
    r.bold = True
    _set_run_font_size(r, size)


def _add_highlighted_bold_run(paragraph, text: str, *, size=_BODY_PT) -> None:
    if not text:
        return
    r = paragraph.add_run(text)
    r.bold = True
    _set_run_font_size(r, size)
    _run_shading_fill(r, _KPI_SHADING_FILL)


def _heading_paragraph(doc: Document, title: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(title.upper())
    run.bold = True
    _set_run_font_size(run, _SECTION_PT)
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)


def _add_achievement_line(doc: Document, row: dict) -> None:
    p = doc.add_paragraph(style=None)
    p.paragraph_format.left_indent = Pt(18)
    p.paragraph_format.first_line_indent = Pt(-18)
    _paragraph_spacing_compact(p)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    verb = row.get("verb", "").strip()
    kpi = row.get("kpi", "").strip()
    rest = " ".join((row.get("rest") or "").split())
    _add_bold_run(p, verb)
    if verb and kpi:
        _add_normal_run(p, " ")
    _add_highlighted_bold_run(p, kpi)
    if rest and kpi:
        _add_normal_run(p, " ")
    elif rest and verb and not kpi:
        _add_normal_run(p, " ")
    _add_normal_run(p, rest)


def _add_experience_bullet(doc: Document, bullet: dict) -> None:
    p = doc.add_paragraph(style=None)
    p.paragraph_format.left_indent = Pt(18)
    p.paragraph_format.first_line_indent = Pt(-18)
    _paragraph_spacing_compact(p)
    verb = bullet.get("verb", "").strip()
    rest = " ".join((bullet.get("rest") or "").split())
    _add_normal_run(p, "• ")
    _add_bold_run(p, verb)
    if verb and rest:
        _add_normal_run(p, " ")
    _add_normal_run(p, rest)


def _add_skill_row(doc: Document, skill: dict) -> None:
    p = doc.add_paragraph()
    _paragraph_spacing_compact(p)
    label = skill.get("label", "").strip()
    value = (skill.get("value") or "").strip()
    _add_bold_run(p, label)
    if label and value:
        _add_normal_run(p, ": ")
    _add_normal_run(p, value)


def build_document(data: dict) -> Document:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Pt(54)
    section.bottom_margin = Pt(54)
    section.left_margin = Pt(54)
    section.right_margin = Pt(54)

    name = data["name"]
    title = data["title"]
    subtitle = data["subtitle"]

    name_p = doc.add_paragraph()
    name_run = name_p.add_run(name)
    name_run.bold = True
    _set_run_font_size(name_run, _NAME_PT)
    name_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _paragraph_spacing_after(name_p, 0)

    tp = doc.add_paragraph()
    tr = tp.add_run(title)
    tr.bold = True
    _set_run_font_size(tr, _TITLE_PT)
    tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _paragraph_spacing_after(tp, 0)

    sub = doc.add_paragraph()
    _add_normal_run(sub, subtitle, size=_BODY_PT)
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _paragraph_spacing_after(sub, 2)

    c = data["contact"]
    contact_line = f'{c["phone"]}  ·  {c["email"]}  ·  {c["linkedin"]}'
    cp = doc.add_paragraph()
    _add_normal_run(cp, contact_line)
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _paragraph_spacing_after(cp, 8)

    _heading_paragraph(doc, "Professional summary")
    sp = doc.add_paragraph()
    summary_text = " ".join((data["summary"] or "").split())
    _add_normal_run(sp, summary_text)

    _heading_paragraph(doc, "Key achievements")
    for achievement in data.get("achievements") or []:
        _add_achievement_line(doc, achievement)

    _heading_paragraph(doc, "Core skills")
    for skill in data.get("skills") or []:
        _add_skill_row(doc, skill)

    _heading_paragraph(doc, "Work experience")
    for job in data.get("experience") or []:
        jp = doc.add_paragraph()
        _paragraph_spacing_after(jp, 2)
        _add_bold_run(jp, job.get("role", "").strip())
        company_p = doc.add_paragraph()
        _paragraph_spacing_compact(company_p)
        meta = f'{job.get("company", "").strip()}  ·  {job.get("period", "").strip()}'
        _add_normal_run(company_p, meta)

        stack = (job.get("stack") or "").strip()
        if stack:
            stp = doc.add_paragraph()
            _paragraph_spacing_compact(stp)
            sr = stp.add_run(stack)
            _set_run_font_size(sr, _BODY_PT)
            sr.italic = True

        for bullet in job.get("bullets") or []:
            _add_experience_bullet(doc, bullet)

    _heading_paragraph(doc, "Education")
    ed = data.get("education") or {}
    ed_p = doc.add_paragraph()
    _add_bold_run(ed_p, ed.get("degree", "").strip())
    uni = doc.add_paragraph()
    _paragraph_spacing_compact(uni)
    _add_normal_run(
        uni, f'{ed.get("university", "").strip()}  ·  {ed.get("period", "").strip()}'
    )

    _heading_paragraph(doc, "Training & courses")
    for course in data.get("courses") or []:
        np = doc.add_paragraph()
        _paragraph_spacing_compact(np)
        _add_bold_run(np, course.get("name", "").strip())
        topics = " ".join((course.get("topics") or "").split())
        if topics:
            tp2 = doc.add_paragraph()
            _paragraph_spacing_compact(tp2)
            _add_normal_run(tp2, topics)

    _heading_paragraph(doc, "Languages")
    for lang in data.get("languages") or []:
        lp = doc.add_paragraph()
        _paragraph_spacing_compact(lp)
        nm = lang.get("name", "").strip()
        level = lang.get("level", "")
        level_text = " ".join(level.split()) if isinstance(level, str) else str(level)
        _add_bold_run(lp, nm)
        _add_normal_run(lp, " — ")
        _add_normal_run(lp, level_text)

    return doc


def main() -> None:
    ap = argparse.ArgumentParser(description="Build English résumé DOCX from resume_data_en.RESUME")
    ap.add_argument(
        "--output",
        "-o",
        type=Path,
        default=_DEFAULT_OUTPUT,
        help=f"Destination .docx (default: {_DEFAULT_OUTPUT})",
    )
    ap.add_argument(
        "--resume-root",
        type=Path,
        default=_DEFAULT_RESUME_ROOT,
        help="Directory containing resume_data_en.py (default: parent of scripts/)",
    )
    args = ap.parse_args()
    output = args.output.expanduser().resolve()

    resume = _load_resume(args.resume_root)
    doc = build_document(resume)
    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output))


if __name__ == "__main__":
    main()
