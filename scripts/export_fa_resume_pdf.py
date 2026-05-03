#!/usr/bin/env python3
"""Rebuild Persian résumé PDF from DOCX (RTL + isolated LTR for English tech runs).

Uses OOXML ``w:rtl`` per run — same convention as Mohammad_Gholampour_Resume_FA.docx.
Fonts are bundled under the repo root ``fonts/`` as static TTF (Noto Sans Arabic + Noto Sans);
CSS ``@font-face`` loads them so the Lunaya Dockerfile does **not** need extra font packages.

Run with Lunaya Python image mounted to this repo (existing image OK; rebuild only when you bump deps elsewhere)::

docker compose -f docker-compose.dev.yml run --rm --no-deps \\
  -v /path/to/mohammadgholampour.github.io:/resume \\
  lunaya-api python /resume/scripts/export_fa_resume_pdf.py \\
    --input /resume/Mohammad_Gholampour_Resume_FA.docx \\
    --output /resume/Mohammad_Gholampour_Resume_FA.pdf

``--resume-root`` overrides the repo root used as WeasyPrint ``base_url`` (defaults to parent of ``scripts/``).
"""

from __future__ import annotations

import argparse
import html
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph
from jinja2 import BaseLoader, Environment, select_autoescape
from weasyprint import HTML

W_MAIN = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"


def run_has_rtl_flag(run) -> bool:
    rpr = run._element.find(W_MAIN + "rPr")
    if rpr is None:
        return False
    rtl_el = rpr.find(qn("w:rtl"))
    if rtl_el is None:
        return False
    val = rtl_el.get(qn("w:val"))
    if val in ("0", "false", "FALSE"):
        return False
    return True


def run_color(run) -> str | None:
    rpr = run._element.find(W_MAIN + "rPr")
    if rpr is None:
        return None
    cel = rpr.find(qn("w:color"))
    if cel is None:
        return None
    val = cel.get(qn("w:val"))
    if not val or val == "auto":
        return None
    hv = "".join(c for c in val.upper() if c in "0123456789ABCDEF")
    return f"#{hv}" if len(hv) == 6 else None


def run_font_half_points(run) -> int | None:
    rpr = run._element.find(W_MAIN + "rPr")
    if rpr is None:
        return None
    sz = rpr.find(qn("w:sz"))
    if sz is None:
        return None
    raw = sz.get(qn("w:val"))
    if not raw:
        return None
    try:
        return int(raw)
    except ValueError:
        return None


def run_is_bold(run) -> bool:
    rpr = run._element.find(W_MAIN + "rPr")
    if rpr is None:
        return False
    el = rpr.find(qn("w:b"))
    if el is None:
        return False
    val = el.get(qn("w:val"))
    if val in ("0", "false", "FALSE"):
        return False
    return True


def run_is_italic(run) -> bool:
    rpr = run._element.find(W_MAIN + "rPr")
    if rpr is None:
        return False
    el = rpr.find(qn("w:i"))
    if el is None:
        return False
    val = el.get(qn("w:val"))
    if val in ("0", "false", "FALSE"):
        return False
    return True


def paragraph_has_bottom_border(p: Paragraph) -> bool:
    ppr = p._element.find(W_MAIN + "pPr")
    if ppr is None:
        return False
    pbdr = ppr.find(qn("w:pBdr"))
    if pbdr is None:
        return False
    return pbdr.find(qn("w:bottom")) is not None


def paragraph_inline_spacing(p: Paragraph) -> str:
    """CSS snippets for margins and line-height from OOXML spacing."""
    ppr = p._element.find(W_MAIN + "pPr")
    if ppr is None:
        return ""
    spo = ppr.find(qn("w:spacing"))
    if spo is None:
        return ""

    bits: list[str] = []

    try:
        b = spo.get(qn("w:before"))
        if b:
            bits.append(f"margin-top:{max(1, round(int(b) / 20)) / 72:.4f}in")
    except (TypeError, ValueError):
        pass
    try:
        a = spo.get(qn("w:after"))
        if a:
            bits.append(f"margin-bottom:{max(1, round(int(a) / 20)) / 72:.4f}in")
    except (TypeError, ValueError):
        pass
    try:
        line = spo.get(qn("w:line"))
        rule = spo.get(qn("w:lineRule"))
        if line and rule == "auto":
            lh = round(int(line)) / 240.0
            bits.append(f"line-height:{lh:.2f}")
    except (TypeError, ValueError):
        pass

    return ";".join(bits)


def paragraph_text(p: Paragraph) -> str:
    return "".join(run.text or "" for run in p.runs)


def max_run_half_points(p: Paragraph) -> int:
    m = 0
    for r in p.runs:
        hp = run_font_half_points(r)
        if hp is not None:
            m = max(m, hp)
    return m


def format_run_markup(run) -> str:
    fragment = run.text if run.text is not None else ""
    if fragment == "":
        return ""

    escaped = html.escape(fragment)
    deco: list[str] = []

    clr = run_color(run)
    if clr:
        deco.append(f"color:{clr}")

    hp = run_font_half_points(run)
    if hp:
        deco.append(f"font-size:{hp / 2.0:.1f}pt")

    inner = escaped
    if deco:
        inner = f'<span style="{html.escape(";".join(deco))}">{inner}</span>'

    if run_is_bold(run):
        inner = f"<strong>{inner}</strong>"
    if run_is_italic(run):
        inner = f"<em>{inner}</em>"

    if run_has_rtl_flag(run):
        return f'<span class="rtl-run" dir="rtl">{inner}</span>'
    return f'<span class="ltr" dir="ltr" style="unicode-bidi:isolate">{inner}</span>'


def paragraph_body_html(p: Paragraph) -> str:
    return "".join(format_run_markup(r) for r in p.runs)


def classify_paragraph(idx: int, first_nonempty_idx: int, p: Paragraph) -> tuple[str, str]:
    """Return (html_tag, extra_class_or_empty_for_section_heading_inner)."""

    style_name = getattr(p.style, "name", None) or ""

    txt = paragraph_text(p).replace("\xa0", " ").strip()
    if not txt:
        return ("skip", "")

    if style_name.startswith("Heading"):
        digits = "".join(c for c in style_name if c.isdigit())
        try:
            n = max(1, min(6, int(digits or "2")))
        except ValueError:
            n = 2
        return (f"h{n}", "")

    if paragraph_has_bottom_border(p):
        return ("h2", "section-heading")

    if idx == first_nonempty_idx and max_run_half_points(p) >= 44 and paragraph_body_html(p):
        return ("h1", "")

    return ("p", "")


TEMPLATE = """<!DOCTYPE html>
<html lang="fa" dir="rtl">
<head>
<meta charset="utf-8"/>
<style>
  @page { size: A4; margin: 12mm 16mm 14mm 16mm; }

  /* Bundled TTF files in repo-root fonts/ — see fonts/OFL-*.txt */
  @font-face {
    font-family: "Noto Sans Arabic";
    src: url("fonts/NotoSansArabic-Regular.ttf") format("truetype");
    font-weight: 400;
    font-style: normal;
    font-display: swap;
  }
  @font-face {
    font-family: "Noto Sans Arabic";
    src: url("fonts/NotoSansArabic-Bold.ttf") format("truetype");
    font-weight: 700;
    font-style: normal;
    font-display: swap;
  }
  @font-face {
    font-family: "Noto Sans";
    src: url("fonts/NotoSans-Regular.ttf") format("truetype");
    font-weight: 400;
    font-style: normal;
    font-display: swap;
  }
  @font-face {
    font-family: "Noto Sans";
    src: url("fonts/NotoSans-Bold.ttf") format("truetype");
    font-weight: 700;
    font-style: normal;
    font-display: swap;
  }

  html, body {
    direction: rtl;
    unicode-bidi: plaintext;
    margin: 0;
    padding: 0;
    color: #333333;
    font-size: 11pt;
    font-family: "Noto Sans Arabic", "Noto Sans", sans-serif;
  }

  span.ltr {
    direction: ltr;
    unicode-bidi: isolate;
    font-family: "Noto Sans", sans-serif;
  }

  span.rtl-run { unicode-bidi: isolate; }

  h1 {
    direction: rtl;
    text-align: right;
    font-size: 20pt;
    font-weight: 700;
    color: #1F3A5F;
    margin: 0 0 0.35em;
  }

  .section-heading {
    direction: rtl;
    text-align: right;
    color: #1F3A5F;
    font-weight: 700;
    font-size: 13pt;
    padding-bottom: 0.06in;
    border-bottom: 1.25pt solid #1F3A5F;
    margin: 0.65em 0 0.4em;
  }

  .section-heading + p { margin-top: 0.35em; }

  h2 {
    direction: rtl;
    text-align: right;
    color: #1F3A5F;
    font-weight: 700;
    font-size: 13pt;
    margin: 0.65em 0 0.4em;
  }

  h2.section-heading-inner {
    border-bottom: none;
    padding-bottom: 0;
    margin: 0;
  }

  p {
    direction: rtl;
    text-align: right;
    text-align-last: right;
    margin: 0 0 0.45em;
    line-height: 1.35;
  }

  code, pre {
    direction: ltr;
    unicode-bidi: isolate;
    font-family: "Noto Sans", monospace;
  }
</style>
</head>
<body>
{% for row in rows %}
{% if row.kind == 'h1' %}
<h1{% if row.pstyle %} style="{{ row.pstyle }}"{% endif %}>{{ row.inner|safe }}</h1>
{% elif row.kind.startswith('h') and row.kind != 'h1' %}
<{{ row.kind }}{% if row.css_class %} class="{{ row.css_class }}"{% endif %}{% if row.pstyle %} style="{{ row.pstyle }}"{% endif %}>{{ row.inner|safe }}</{{ row.kind }}>
{% else %}
<p{% if row.pstyle %} style="{{ row.pstyle }}"{% endif %}>{{ row.inner|safe }}</p>
{% endif %}
{% endfor %}
</body>
</html>
"""


def build_rows(doc: Document) -> list[dict]:
    first_nonempty = None
    for i, p in enumerate(doc.paragraphs):
        if paragraph_text(p).strip():
            first_nonempty = i
            break
    if first_nonempty is None:
        return []

    out: list[dict] = []
    for idx, p in enumerate(doc.paragraphs):
        tag_or_skip, klass = classify_paragraph(idx, first_nonempty, p)
        if tag_or_skip == "skip":
            continue
        inner = paragraph_body_html(p)
        if not inner.strip():
            continue
        pst = paragraph_inline_spacing(p)

        if klass == "section-heading":
            out.append(dict(kind="h2", inner=inner, css_class="section-heading", pstyle=pst))
            continue

        if tag_or_skip == "p":
            out.append(dict(kind="p", inner=inner, css_class="", pstyle=pst))
        else:
            out.append(dict(kind=tag_or_skip, inner=inner, css_class="", pstyle=pst))

    return out


def main() -> None:
    parser = argparse.ArgumentParser(description="Export FA résumé DOCX → RTL PDF via WeasyPrint.")
    parser.add_argument("--input", type=Path, required=True)
    parser.add_argument("--output", type=Path, required=True)
    parser.add_argument(
        "--resume-root",
        type=Path,
        default=None,
        help="Repo root containing fonts/ (default: parent of scripts/). Used as WeasyPrint base URL.",
    )
    args = parser.parse_args()

    resume_root = (args.resume_root or Path(__file__).resolve().parent.parent).resolve()
    fonts_dir = resume_root / "fonts"
    required_ttfs = (
        fonts_dir / "NotoSansArabic-Regular.ttf",
        fonts_dir / "NotoSansArabic-Bold.ttf",
        fonts_dir / "NotoSans-Regular.ttf",
        fonts_dir / "NotoSans-Bold.ttf",
    )
    missing = [str(p.relative_to(resume_root)) for p in required_ttfs if not p.is_file()]
    if missing:
        raise SystemExit("Missing bundled font files under resume root: " + ", ".join(missing))

    doc = Document(str(args.input))
    rows = build_rows(doc)

    env = Environment(loader=BaseLoader(), autoescape=select_autoescape(["html", "xml"]))
    html_doc = env.from_string(TEMPLATE).render(rows=rows)

    args.output.parent.mkdir(parents=True, exist_ok=True)

    base_uri = resume_root.as_uri().rstrip("/") + "/"
    HTML(string=html_doc, base_url=base_uri).write_pdf(target=str(args.output))


if __name__ == "__main__":
    main()
